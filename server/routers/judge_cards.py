import base64
import hashlib
import hmac
import io
import json
import os
import re
from datetime import datetime, timedelta, timezone
from urllib.parse import urlencode, urljoin

import qrcode
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from sqlmodel import SQLModel, Session, select

from access import require_competition_access
from auth import require_staff
from database import get_session
from models import (
    CompetitionHeat,
    CompetitionHeatAssignment,
    CompetitionParticipant,
    CompetitionPhase,
    Participant,
)

router = APIRouter(prefix="/api/judge-cards", tags=["judge_cards"])

CARD_INNER_PADDING = 6
CARD_QR_FIXED_SIZE_INCHES = 0.72
CARD_QR_FIXED_SIZE_PT = CARD_QR_FIXED_SIZE_INCHES * 72
CARD_QR_SAFE_ZONE_PT = 8
CARD_TITLE_BAND_PT = 20
CARD_LEFT_COLUMN_MIN_WIDTH_PT = 92
CARD_COLUMN_GAP_PT = 12


class JudgeCardsExportBody(SQLModel):
    competition_id: int
    phase_ids: list[int] | None = None
    categories: list[str] | None = None
    only_confirmed: int = 1
    include_unassigned: int = 1
    sort_mode: str = "phase_heat_lane_name"
    layout: str = "auto"
    include_score_field: int = 1
    include_signature_field: int = 1
    include_notes_field: int = 0
    include_qr: int = 1
    extra_fields: list[str] | None = None
    title: str | None = None
    judge_portal_base_url: str | None = None
    judge_portal_path: str = "judge/score"
    qr_expiration_days: int = 30
    page_size: str = "letter"
    font_scale: float = 1.0
    line_spacing: float = 1.0
    writing_space_chars: int = 30


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _b64url_encode(raw: bytes) -> str:
    return base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")


def _qr_secret() -> str:
    value = (os.getenv("CHECKIN_QR_SECRET") or os.getenv("SECRET_KEY") or "").strip()
    if not value:
        raise HTTPException(500, "Falta CHECKIN_QR_SECRET o SECRET_KEY en el servidor")
    return value


def _make_judge_token(
    *,
    competition_id: int,
    user_id: int,
    phase_id: int,
    heat_id: int | None,
    expires_days: int,
) -> str:
    now = _utcnow()
    safe_days = min(max(int(expires_days or 30), 1), 365)
    payload = {
        "scope": "judge_score",
        "c": int(competition_id),
        "p": int(user_id),
        "ph": int(phase_id),
        "h": int(heat_id) if heat_id else None,
        "iat": int(now.timestamp()),
        "exp": int((now + timedelta(days=safe_days)).timestamp()),
    }
    payload_b64 = _b64url_encode(json.dumps(payload, separators=(",", ":")).encode("utf-8"))
    signature = hmac.new(_qr_secret().encode("utf-8"), payload_b64.encode("utf-8"), hashlib.sha256).digest()
    return f"{payload_b64}.{_b64url_encode(signature)}"


def _qr_image_bytes(url: str) -> io.BytesIO:
    qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=5, border=1)
    qr.add_data(url)
    qr.make(fit=True)
    image = qr.make_image(fill_color="#0D0F12", back_color="#F5F7FA")
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _visible_card_line_budget(
    *,
    include_score_field: bool,
    include_signature_field: bool,
    include_notes_field: bool,
    include_qr: bool,
    extra_fields: set[str],
) -> int:
    lines = 2  # name + base metadata
    if "cedula" in extra_fields:
        lines += 1
    if include_score_field:
        lines += 1
    if include_signature_field:
        lines += 1
    if include_notes_field:
        lines += 1
    if include_qr:
        lines += 1
    return lines


def _safe_font_scale(value: float | int | None) -> float:
    try:
        return min(max(float(value or 1.0), 0.75), 1.35)
    except (TypeError, ValueError):
        return 1.0


def _safe_line_spacing(value: float | int | None) -> float:
    try:
        return min(max(float(value or 1.0), 0.8), 1.8)
    except (TypeError, ValueError):
        return 1.0


def _safe_writing_space_chars(value: int | float | None) -> int:
    try:
        return min(max(int(value or 30), 8), 48)
    except (TypeError, ValueError):
        return 30


def _writing_field_text(label: str, writing_space_chars: int) -> str:
    return f"{label}: {'_' * _safe_writing_space_chars(writing_space_chars)}"


def _resolve_layout_name(
    layout: str,
    *,
    include_score_field: bool,
    include_signature_field: bool,
    include_notes_field: bool,
    include_qr: bool,
    extra_fields: set[str],
) -> str:
    normalized = (layout or "").strip().lower()
    custom_match = re.fullmatch(r"([1-4])x([1-9]|10)", normalized)
    allowed = {"auto", "2x3", "2x4", "2x5", "3x3", "3x4", "3x5", "3x6"}
    if custom_match:
        return normalized
    if normalized not in allowed:
        normalized = "auto"
    if normalized != "auto":
        return normalized
    visible_lines = _visible_card_line_budget(
        include_score_field=include_score_field,
        include_signature_field=include_signature_field,
        include_notes_field=include_notes_field,
        include_qr=include_qr,
        extra_fields=extra_fields,
    )
    if visible_lines <= 4:
        return "3x6"
    if visible_lines <= 5:
        return "3x5"
    if visible_lines <= 6:
        return "3x4"
    if visible_lines <= 7:
        return "2x5"
    if visible_lines <= 8:
        return "2x4"
    return "2x3"


def _layout_to_grid(layout: str) -> tuple[int, int]:
    normalized = (layout or "").strip().lower()
    cols, rows = normalized.split("x", 1)
    return int(cols), int(rows)


def _set_cell_border(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_tag = f"w:{edge}"
        element = tc_borders.find(qn(edge_tag))
        if element is None:
            element = OxmlElement(edge_tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "252A33")


def _set_page_style(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)


def _add_header(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x0D, 0x0F, 0x12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _build_cards_dataset(
    session: Session,
    *,
    competition_id: int,
    phase_ids: list[int] | None,
    categories: list[str] | None,
    only_confirmed: bool,
    include_unassigned: bool,
) -> tuple[list[dict], dict[int, CompetitionPhase]]:
    phase_query = select(CompetitionPhase).where(CompetitionPhase.competition_id == competition_id)
    if phase_ids:
        phase_query = phase_query.where(CompetitionPhase.id.in_(phase_ids))
    phase_query = phase_query.order_by(CompetitionPhase.block_order, CompetitionPhase.orden, CompetitionPhase.id)
    phases = session.exec(phase_query).all()
    if not phases:
        raise HTTPException(404, "No se encontraron fases para generar tarjetas")
    phase_map = {int(p.id): p for p in phases if p.id is not None}

    cp_query = (
        select(CompetitionParticipant, Participant)
        .join(Participant, Participant.id == CompetitionParticipant.user_id)
        .where(CompetitionParticipant.competition_id == competition_id)
    )
    if only_confirmed:
        cp_query = cp_query.where(CompetitionParticipant.estado == "confirmado")
    cp_rows = session.exec(cp_query).all()

    category_filter = {str(c or "").strip().lower() for c in (categories or []) if str(c or "").strip()}
    participants_pool: dict[int, dict] = {}
    for cp, participant in cp_rows:
        category = str(cp.categoria or participant.categoria or "").strip()
        if category_filter and category.lower() not in category_filter:
            continue
        participants_pool[int(participant.id)] = {
            "user_id": int(participant.id),
            "participant_name": f"{(participant.nombre or '').strip()} {(participant.apellido or '').strip()}".strip(),
            "cedula": str(participant.cedula or "").strip(),
            "category": category or "Sin categoria",
        }
    if not participants_pool:
        raise HTTPException(404, "No hay participantes que coincidan con los filtros")

    heat_query = (
        select(CompetitionHeatAssignment, CompetitionHeat)
        .join(CompetitionHeat, CompetitionHeat.id == CompetitionHeatAssignment.heat_id)
        .where(CompetitionHeat.competition_id == competition_id)
        .where(CompetitionHeat.phase_id.in_(list(phase_map.keys())))
        .order_by(CompetitionHeat.phase_id, CompetitionHeat.heat_number, CompetitionHeat.start_at, CompetitionHeatAssignment.lane_number)
    )
    heat_rows = session.exec(heat_query).all()

    assigned_by_phase: dict[int, list[dict]] = {}
    for assignment, heat in heat_rows:
        user_id = int(assignment.user_id or 0)
        if user_id <= 0:
            continue
        base = participants_pool.get(user_id)
        if not base:
            continue
        pid = int(heat.phase_id)
        assigned_by_phase.setdefault(pid, []).append(
            {
                **base,
                "phase_id": pid,
                "heat_id": int(heat.id),
                "phase_name": str(phase_map[pid].nombre or "").strip() or f"Fase {pid}",
                "block_name": str(phase_map[pid].block_name or "").strip(),
                "heat_name": str(heat.nombre or "").strip() or f"Heat {int(heat.heat_number or 0)}",
                "heat_number": int(heat.heat_number or 0),
                "lane_number": int(assignment.lane_number or 0),
                "start_at": heat.start_at,
                "location_name": str(heat.location_name or "").strip(),
            }
        )

    cards: list[dict] = []
    for phase_id, phase in phase_map.items():
        phase_cards = assigned_by_phase.get(phase_id, [])
        if phase_cards:
            cards.extend(phase_cards)
            continue
        if include_unassigned:
            for item in participants_pool.values():
                cards.append(
                    {
                        **item,
                        "phase_id": int(phase_id),
                        "heat_id": None,
                        "phase_name": str(phase.nombre or "").strip() or f"Fase {phase_id}",
                        "block_name": str(phase.block_name or "").strip(),
                        "heat_name": "",
                        "heat_number": 0,
                        "lane_number": 0,
                        "start_at": None,
                        "location_name": "",
                    }
                )
    if not cards:
        raise HTTPException(404, "No se encontraron tarjetas para exportar con los filtros actuales")
    return cards, phase_map


def _sort_cards(cards: list[dict], sort_mode: str) -> list[dict]:
    mode = (sort_mode or "").strip().lower()
    if mode == "name":
        return sorted(cards, key=lambda c: (c["participant_name"].lower(), int(c["phase_id"]), int(c.get("lane_number") or 0)))
    return sorted(
        cards,
        key=lambda c: (
            int(c["phase_id"]),
            int(c.get("heat_number") or 0),
            int(c.get("lane_number") or 0),
            c["participant_name"].lower(),
        ),
    )


def _card_meta_lines(*, card: dict, include_cedula: bool) -> list[tuple[str, float, str]]:
    lines: list[tuple[str, float, str]] = [
        (f"{card.get('phase_name', '')} | {card.get('category', '')}", 6.7, "#252A33")
    ]
    heat_parts = []
    if card.get("heat_name"):
        heat_parts.append(f"Heat: {card['heat_name']}")
    if int(card.get("lane_number") or 0) > 0:
        heat_parts.append(f"Carril: {int(card['lane_number'])}")
    if card.get("location_name"):
        heat_parts.append(f"Zona: {card['location_name']}")
    if heat_parts:
        lines.append((" | ".join(heat_parts), 5.8, "#6B7280"))
    if include_cedula and card.get("cedula"):
        lines.append((f"ID: {card['cedula']}", 5.8, "#6B7280"))
    return lines


def _card_form_lines(*, include_score_field: bool, include_signature_field: bool, include_notes_field: bool, writing_space_chars: int) -> list[tuple[str, float, str]]:
    lines: list[tuple[str, float, str]] = []
    if include_score_field:
        lines.append((_writing_field_text("Puntuacion", writing_space_chars), 7.5, "#0D0F12"))
    if include_signature_field:
        lines.append((_writing_field_text("Firma atleta", writing_space_chars), 7.5, "#0D0F12"))
    if include_notes_field:
        lines.append((_writing_field_text("Notas", writing_space_chars), 6.5, "#6B7280"))
    return lines


def _card_column_lines(*, card: dict, include_score_field: bool, include_signature_field: bool, include_notes_field: bool, include_cedula: bool, writing_space_chars: int) -> list[tuple[str, float, str, bool]]:
    lines: list[tuple[str, float, str, bool]] = []
    for text, size, color in _card_meta_lines(card=card, include_cedula=include_cedula):
        lines.append((text, size, color, False))
    for text, size, color in _card_form_lines(
        include_score_field=include_score_field,
        include_signature_field=include_signature_field,
        include_notes_field=include_notes_field,
        writing_space_chars=writing_space_chars,
    ):
        lines.append((text, size, color, color == "#0D0F12"))
    return lines


def _ensure_fixed_qr_fits(*, width: float, height: float, include_qr: bool) -> None:
    if not include_qr:
        return
    right_column_width = CARD_QR_FIXED_SIZE_PT + (CARD_QR_SAFE_ZONE_PT * 2)
    left_column_width = width - (CARD_INNER_PADDING * 2) - right_column_width - CARD_COLUMN_GAP_PT
    body_height = height - CARD_TITLE_BAND_PT - (CARD_INNER_PADDING * 2)
    if left_column_width < CARD_LEFT_COLUMN_MIN_WIDTH_PT or body_height < CARD_QR_FIXED_SIZE_PT:
        raise HTTPException(
            400,
            "La configuracion actual deja las tarjetas demasiado pequenas para mantener un QR fijo y legible. Reduce tarjetas por hoja.",
        )


def _add_card_to_cell(
    cell,
    *,
    card: dict,
    include_score_field: bool,
    include_signature_field: bool,
    include_notes_field: bool,
    include_qr: bool,
    qr_url: str | None,
    extra_fields: set[str],
    font_scale: float,
    line_spacing: float,
    writing_space_chars: int,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_border(cell)
    cell.text = ""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    p_name = cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(0)
    run_name = p_name.add_run(card["participant_name"] or "Participante")
    run_name.bold = True
    run_name.font.size = Pt(10 * _safe_font_scale(font_scale))
    run_name.font.color.rgb = RGBColor(0x0D, 0x0F, 0x12)

    include_cedula = "cedula" in extra_fields
    inner = cell.add_table(rows=1, cols=2)
    inner.autofit = False
    left_cell = inner.cell(0, 0)
    right_cell = inner.cell(0, 1)
    right_width_inches = CARD_QR_FIXED_SIZE_INCHES + ((CARD_QR_SAFE_ZONE_PT * 2) / 72)
    right_cell.width = int(Inches(right_width_inches)) if cell.width else None
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    column_lines = _card_column_lines(
        card=card,
        include_score_field=include_score_field,
        include_signature_field=include_signature_field,
        include_notes_field=include_notes_field,
        include_cedula=include_cedula,
        writing_space_chars=writing_space_chars,
    )
    for text, size, color, bold in column_lines:
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt((size + 1.2) * _safe_font_scale(font_scale) * _safe_line_spacing(line_spacing))
        run = p.add_run(text)
        run.font.size = Pt((size + 0.4) * _safe_font_scale(font_scale))
        run.bold = bool(bold)
        rgb = {"#0D0F12": RGBColor(0x0D, 0x0F, 0x12), "#252A33": RGBColor(0x25, 0x2A, 0x33), "#6B7280": RGBColor(0x6B, 0x72, 0x80)}[color]
        run.font.color.rgb = rgb

    if include_qr and qr_url:
        p_qr = right_cell.add_paragraph()
        p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_qr.paragraph_format.space_before = Pt(0)
        p_qr.paragraph_format.space_after = Pt(0)
        run_qr = p_qr.add_run()
        run_qr.add_picture(_qr_image_bytes(qr_url), width=Inches(CARD_QR_FIXED_SIZE_INCHES))


def _build_portal_url(body: JudgeCardsExportBody) -> str:
    portal_base = (body.judge_portal_base_url or os.getenv("LEADERBOARD_BASE_URL") or "http://localhost:5173/").strip()
    portal_path = (body.judge_portal_path or "judge/score").lstrip("/")
    portal_base = portal_base if portal_base.endswith("/") else f"{portal_base}/"
    return urljoin(portal_base, portal_path)


def _card_qr_url(*, body: JudgeCardsExportBody, card: dict, portal_url: str) -> str:
    token = _make_judge_token(
        competition_id=body.competition_id,
        user_id=int(card["user_id"]),
        phase_id=int(card["phase_id"]),
        heat_id=(int(card["heat_id"]) if card.get("heat_id") else None),
        expires_days=int(body.qr_expiration_days or 30),
    )
    # Keep only token in query to reduce QR density and improve scan reliability.
    query = urlencode({"token": token})
    return f"{portal_url}?{query}"


def _resolve_page_size(raw: str | None) -> tuple[tuple[float, float], str]:
    normalized = str(raw or "letter").strip().lower()
    if normalized in {"a4"}:
        return A4, "a4"
    return LETTER, "letter"


def _draw_pdf_card(
    pdf: canvas.Canvas,
    *,
    x: float,
    y: float,
    width: float,
    height: float,
    card: dict,
    include_score_field: bool,
    include_signature_field: bool,
    include_notes_field: bool,
    include_qr: bool,
    qr_url: str | None,
    include_cedula: bool,
    font_scale: float,
    line_spacing: float,
    writing_space_chars: int,
) -> None:
    pad = CARD_INNER_PADDING
    _ensure_fixed_qr_fits(width=width, height=height, include_qr=include_qr and bool(qr_url))
    pdf.setStrokeColor(HexColor("#252A33"))
    pdf.setLineWidth(0.8)
    pdf.rect(x, y, width, height, stroke=1, fill=0)

    top_band_height = CARD_TITLE_BAND_PT
    name_y = y + height - pad - 8
    pdf.setFillColor(HexColor("#0D0F12"))
    safe_font_scale = _safe_font_scale(font_scale)
    pdf.setFont("Helvetica-Bold", 8.8 * safe_font_scale)
    pdf.drawCentredString(x + (width / 2), name_y, (card.get("participant_name") or "Participante")[:80])

    body_top = y + height - top_band_height - pad
    body_bottom = y + pad
    body_height = body_top - body_bottom
    left_x = x + pad
    qr_column_width = (CARD_QR_FIXED_SIZE_PT + (CARD_QR_SAFE_ZONE_PT * 2)) if include_qr and qr_url else 0
    qr_size = CARD_QR_FIXED_SIZE_PT if include_qr and qr_url else 0
    left_width = width - (pad * 2) - qr_column_width - (CARD_COLUMN_GAP_PT if qr_column_width else 0)
    text_width = max(0, left_width)
    column_lines = _card_column_lines(
        card=card,
        include_score_field=include_score_field,
        include_signature_field=include_signature_field,
        include_notes_field=include_notes_field,
        include_cedula=include_cedula,
        writing_space_chars=writing_space_chars,
    )
    line_step = 7.2 * safe_font_scale * _safe_line_spacing(line_spacing)
    block_height = (max(0, len(column_lines) - 1) * line_step) + ((column_lines[-1][1] * safe_font_scale) if column_lines else 0)
    line_y = body_bottom + ((body_height + block_height) / 2) - 2
    for text, size, color, bold in column_lines:
        pdf.setFillColor(HexColor(color))
        sized_font = size * safe_font_scale
        pdf.setFont("Helvetica-Bold" if bold else "Helvetica", sized_font)
        pdf.drawString(left_x, line_y, text[: max(10, int(text_width / max(sized_font * 0.43, 1)))])
        line_y -= line_step
    if include_qr and qr_url:
        qr_x = x + width - pad - CARD_QR_SAFE_ZONE_PT - qr_size
        qr_y = body_bottom + ((body_height - qr_size) / 2)
        qr_reader = ImageReader(_qr_image_bytes(qr_url))
        pdf.drawImage(qr_reader, qr_x, qr_y, width=qr_size, height=qr_size, preserveAspectRatio=True, mask="auto")


@router.post("/export-docx")
def export_judge_cards_docx(
    body: JudgeCardsExportBody,
    session: Session = Depends(get_session),
    user=Depends(require_staff),
):
    competition = require_competition_access(session, body.competition_id, user)
    cards, _phase_map = _build_cards_dataset(
        session,
        competition_id=body.competition_id,
        phase_ids=body.phase_ids,
        categories=body.categories,
        only_confirmed=bool(body.only_confirmed),
        include_unassigned=bool(body.include_unassigned),
    )
    cards = _sort_cards(cards, body.sort_mode)
    extra_fields = {str(name or "").strip().lower() for name in (body.extra_fields or []) if str(name or "").strip()}
    resolved_layout = _resolve_layout_name(
        body.layout,
        include_score_field=bool(body.include_score_field),
        include_signature_field=bool(body.include_signature_field),
        include_notes_field=bool(body.include_notes_field),
        include_qr=bool(body.include_qr),
        extra_fields=extra_fields,
    )
    cols, rows = _layout_to_grid(resolved_layout)
    cards_per_page = cols * rows
    portal_url = _build_portal_url(body)

    doc = Document()
    _set_page_style(doc)
    title = (body.title or f"Tarjetas de puntuacion - {competition.nombre}").strip()
    subtitle = f"Competencia: {competition.nombre} | Tarjetas: {len(cards)} | Generado: {_utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
    _add_header(doc, title, subtitle)

    for start in range(0, len(cards), cards_per_page):
        chunk = cards[start:start + cards_per_page]
        if start > 0:
            doc.add_page_break()
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = False
        page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        col_width = int(page_width / cols)
        row_height = Inches(9.68 / max(rows, 1))
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = row_height
        for col_idx in range(cols):
            for row_idx in range(rows):
                table.cell(row_idx, col_idx).width = col_width

        for idx, card in enumerate(chunk):
            r = idx // cols
            c = idx % cols
            qr_url = _card_qr_url(body=body, card=card, portal_url=portal_url)
            _add_card_to_cell(
                table.cell(r, c),
                card=card,
                include_score_field=bool(body.include_score_field),
                include_signature_field=bool(body.include_signature_field),
                include_notes_field=bool(body.include_notes_field),
                include_qr=bool(body.include_qr),
                qr_url=qr_url,
                extra_fields=extra_fields,
                font_scale=body.font_scale,
                line_spacing=body.line_spacing,
                writing_space_chars=body.writing_space_chars,
            )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"finalrep_tarjetas_competencia_{int(body.competition_id)}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.post("/export-pdf")
def export_judge_cards_pdf(
    body: JudgeCardsExportBody,
    session: Session = Depends(get_session),
    user=Depends(require_staff),
):
    competition = require_competition_access(session, body.competition_id, user)
    cards, _phase_map = _build_cards_dataset(
        session,
        competition_id=body.competition_id,
        phase_ids=body.phase_ids,
        categories=body.categories,
        only_confirmed=bool(body.only_confirmed),
        include_unassigned=bool(body.include_unassigned),
    )
    cards = _sort_cards(cards, body.sort_mode)
    extra_fields = {str(name or "").strip().lower() for name in (body.extra_fields or []) if str(name or "").strip()}
    resolved_layout = _resolve_layout_name(
        body.layout,
        include_score_field=bool(body.include_score_field),
        include_signature_field=bool(body.include_signature_field),
        include_notes_field=bool(body.include_notes_field),
        include_qr=bool(body.include_qr),
        extra_fields=extra_fields,
    )
    cols, rows = _layout_to_grid(resolved_layout)
    cards_per_page = cols * rows
    page_size, page_size_label = _resolve_page_size(body.page_size)
    page_width, page_height = page_size
    margin = 24
    header_space = 28
    grid_top = page_height - margin - header_space
    grid_width = page_width - (margin * 2)
    grid_height = page_height - (margin * 2) - header_space
    cell_w = grid_width / cols
    cell_h = grid_height / rows
    include_cedula = "cedula" in extra_fields
    portal_url = _build_portal_url(body)
    title = (body.title or f"Tarjetas de puntuacion - {competition.nombre}").strip()

    buf = io.BytesIO()
    pdf = canvas.Canvas(buf, pagesize=page_size)
    total_pages = max(1, (len(cards) + cards_per_page - 1) // cards_per_page)

    for page_index, start in enumerate(range(0, len(cards), cards_per_page), 1):
        chunk = cards[start:start + cards_per_page]
        pdf.setFont("Helvetica-Bold", 11)
        pdf.setFillColor(HexColor("#0D0F12"))
        pdf.drawString(margin, page_height - margin + 4, title[:100])
        pdf.setFont("Helvetica", 7)
        pdf.setFillColor(HexColor("#6B7280"))
        stamp = _utcnow().strftime("%Y-%m-%d %H:%M UTC")
        pdf.drawString(margin, page_height - margin - 8, f"Competencia: {competition.nombre} | Pagina {page_index}/{total_pages} | {page_size_label.upper()} | {stamp}")

        for idx, card in enumerate(chunk):
            row = idx // cols
            col = idx % cols
            x = margin + (col * cell_w)
            y = grid_top - ((row + 1) * cell_h)
            qr_url = _card_qr_url(body=body, card=card, portal_url=portal_url)
            _draw_pdf_card(
                pdf,
                x=x,
                y=y,
                width=cell_w,
                height=cell_h,
                card=card,
                include_score_field=bool(body.include_score_field),
                include_signature_field=bool(body.include_signature_field),
                include_notes_field=bool(body.include_notes_field),
                include_qr=bool(body.include_qr),
                qr_url=qr_url,
                include_cedula=include_cedula,
                font_scale=body.font_scale,
                line_spacing=body.line_spacing,
                writing_space_chars=body.writing_space_chars,
            )
        if page_index < total_pages:
            pdf.showPage()

    pdf.save()
    buf.seek(0)
    filename = f"finalrep_tarjetas_competencia_{int(body.competition_id)}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
